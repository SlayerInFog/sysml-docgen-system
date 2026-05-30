from collections import Counter
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from jinja2 import Template
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy.orm import Session

from app.models.document import GeneratedDocument
from app.models.sysml import ModelElement, ModelRelation, SysMLModel


DEFAULT_TEMPLATE = """<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>{{ title }}</title>
  <style>
    body { font-family: "Microsoft YaHei", sans-serif; line-height: 1.75; color: #1f2937; }
    h1, h2, h3 { color: #0f172a; }
    table { border-collapse: collapse; width: 100%; margin: 16px 0; }
    th, td { border: 1px solid #cbd5e1; padding: 8px 10px; text-align: left; }
    th { background: #e2e8f0; }
    .meta { color: #64748b; }
  </style>
</head>
<body>
  <h1>{{ title }}</h1>
  <p class="meta">模型名称：{{ model.name }} ｜ 模型版本：V{{ model.version }}</p>

  <h2>1. 模型概述</h2>
  <p>本文件由 SysML 模型自动生成，覆盖模型元素、结构关系和需求/块/活动等关键内容。</p>
  <ul>
    <li>元素总数：{{ stats.total_elements }}</li>
    <li>关系总数：{{ stats.total_relations }}</li>
    <li>主要元素类型：{{ stats.top_types }}</li>
  </ul>

  <h2>2. 模型元素清单</h2>
  <table>
    <thead><tr><th>编号</th><th>名称</th><th>类型</th><th>说明</th></tr></thead>
    <tbody>
    {% for element in elements %}
      <tr>
        <td>{{ loop.index }}</td>
        <td>{{ element.name }}</td>
        <td>{{ element.type }}</td>
        <td>{{ element.documentation or "" }}</td>
      </tr>
    {% endfor %}
    </tbody>
  </table>

  <h2>3. 模型关系</h2>
  <table>
    <thead><tr><th>源元素</th><th>关系</th><th>目标元素</th></tr></thead>
    <tbody>
    {% for relation in relations %}
      <tr>
        <td>{{ element_names.get(relation.source_uid, relation.source_uid) }}</td>
        <td>{{ relation.relation_type }}</td>
        <td>{{ element_names.get(relation.target_uid, relation.target_uid) }}</td>
      </tr>
    {% endfor %}
    </tbody>
  </table>

  <h2>4. 结论</h2>
  <p>该文档体现了模型驱动文档生成思路，模型数据作为单一可信源，文档内容可随模型更新重新生成。</p>
</body>
</html>"""


MAX_OFFICE_TABLE_ROWS = 300


# 用模板和模型数据渲染 HTML 文档。
def render_document_html(
    db: Session,
    model: SysMLModel,
    title: str,
    template_content: str | None = None,
) -> str:
    elements = db.query(ModelElement).filter(ModelElement.model_id == model.id).order_by(ModelElement.type, ModelElement.name).all()
    relations = db.query(ModelRelation).filter(ModelRelation.model_id == model.id).all()
    type_counts = Counter(element.type for element in elements)
    element_names = {element.element_uid: element.name for element in elements}

    context = {
        "title": title,
        "model": model,
        "elements": elements,
        "relations": relations,
        "element_names": element_names,
        "stats": {
            "total_elements": len(elements),
            "total_relations": len(relations),
            "top_types": "、".join(f"{name}({count})" for name, count in type_counts.most_common(5)) or "无",
        },
    }
    return Template(template_content or DEFAULT_TEMPLATE).render(**context)


# 将 HTML 内容转换为 DOCX 文件。
def html_to_docx(html: str, output_path: Path) -> None:
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    title = soup.find("h1")
    if title:
        doc.add_heading(title.get_text(strip=True), level=0)
    for node in soup.find_all(["h2", "h3", "p", "li"]):
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        if node.name == "h2":
            doc.add_heading(text, level=1)
        elif node.name == "h3":
            doc.add_heading(text, level=2)
        elif node.name == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue
        exported_rows = rows[:MAX_OFFICE_TABLE_ROWS]
        max_cols = max(len(row.find_all(["th", "td"])) for row in exported_rows)
        doc_table = doc.add_table(rows=len(exported_rows), cols=max_cols)
        doc_table.style = "Table Grid"
        for r_idx, row in enumerate(exported_rows):
            cells = row.find_all(["th", "td"])
            for c_idx, cell in enumerate(cells):
                doc_table.cell(r_idx, c_idx).text = cell.get_text(" ", strip=True)
        if len(rows) > MAX_OFFICE_TABLE_ROWS:
            doc.add_paragraph(
                f"表格内容较多，DOCX 导出仅展示前 {MAX_OFFICE_TABLE_ROWS} 行；完整内容请使用 HTML 导出。"
            )
    doc.save(output_path)


# 将 HTML 内容转换为 PDF 文件。
def html_to_pdf(html: str, output_path: Path) -> None:
    soup = BeautifulSoup(html, "html.parser")
    story = []
    styles = getSampleStyleSheet()
    for node in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        style = styles["BodyText"]
        if node.name == "h1":
            style = styles["Title"]
        elif node.name == "h2":
            style = styles["Heading1"]
        elif node.name == "h3":
            style = styles["Heading2"]
        elif node.name == "li":
            text = "- " + text
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 8))
    SimpleDocTemplate(str(output_path), pagesize=A4).build(story)


# 处理 export_model_docx_summary 相关逻辑。
def export_model_docx_summary(db: Session, doc_record: GeneratedDocument, output_path: Path, max_rows: int = 300) -> None:
    elements = (
        db.query(ModelElement)
        .filter(ModelElement.model_id == doc_record.model_id)
        .order_by(ModelElement.type, ModelElement.name)
        .limit(max_rows)
        .all()
    )
    relations = db.query(ModelRelation).filter(ModelRelation.model_id == doc_record.model_id).limit(max_rows).all()
    total_elements = db.query(ModelElement).filter(ModelElement.model_id == doc_record.model_id).count()
    total_relations = db.query(ModelRelation).filter(ModelRelation.model_id == doc_record.model_id).count()

    document = Document()
    document.add_heading(doc_record.title, level=0)
    document.add_paragraph(f"模型名称：{doc_record.model.name}")
    document.add_paragraph(f"模型版本：V{doc_record.model.version}")
    document.add_paragraph(f"元素总数：{total_elements}")
    document.add_paragraph(f"关系总数：{total_relations}")
    if total_elements > max_rows or total_relations > max_rows:
        document.add_paragraph(f"模型规模较大，DOCX 导出展示前 {max_rows} 条元素和关系；完整内容请导出 HTML。")

    document.add_heading("模型元素", level=1)
    element_table = document.add_table(rows=1, cols=4)
    element_table.style = "Table Grid"
    for idx, title in enumerate(["序号", "名称", "类型", "说明"]):
        element_table.cell(0, idx).text = title
    for idx, element in enumerate(elements, start=1):
        row = element_table.add_row().cells
        row[0].text = str(idx)
        row[1].text = element.name or ""
        row[2].text = element.type or ""
        row[3].text = element.documentation or ""

    element_names = {item.element_uid: item.name for item in elements}
    document.add_heading("模型关系", level=1)
    relation_table = document.add_table(rows=1, cols=4)
    relation_table.style = "Table Grid"
    for idx, title in enumerate(["序号", "源元素", "关系", "目标元素"]):
        relation_table.cell(0, idx).text = title
    for idx, relation in enumerate(relations, start=1):
        row = relation_table.add_row().cells
        row[0].text = str(idx)
        row[1].text = element_names.get(relation.source_uid, relation.source_uid)
        row[2].text = relation.relation_type or ""
        row[3].text = element_names.get(relation.target_uid, relation.target_uid)
    document.save(output_path)


# 处理 export_model_pdf_summary 相关逻辑。
def export_model_pdf_summary(db: Session, doc_record: GeneratedDocument, output_path: Path, max_rows: int = 120) -> None:
    elements = (
        db.query(ModelElement)
        .filter(ModelElement.model_id == doc_record.model_id)
        .order_by(ModelElement.type, ModelElement.name)
        .limit(max_rows)
        .all()
    )
    total_elements = db.query(ModelElement).filter(ModelElement.model_id == doc_record.model_id).count()
    total_relations = db.query(ModelRelation).filter(ModelRelation.model_id == doc_record.model_id).count()

    styles = getSampleStyleSheet()
    story = [
        Paragraph(doc_record.title, styles["Title"]),
        Spacer(1, 10),
        Paragraph(f"Model: {doc_record.model.name} V{doc_record.model.version}", styles["BodyText"]),
        Paragraph(f"Elements: {total_elements}", styles["BodyText"]),
        Paragraph(f"Relations: {total_relations}", styles["BodyText"]),
        Paragraph(f"Large model export preview: first {max_rows} elements. Use HTML for full content.", styles["BodyText"]),
        Spacer(1, 10),
        Paragraph("Model Elements", styles["Heading1"]),
    ]
    for idx, element in enumerate(elements, start=1):
        text = f"{idx}. {element.name} [{element.type}]"
        if element.documentation:
            text += f" - {element.documentation[:200]}"
        story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 5))
    SimpleDocTemplate(str(output_path), pagesize=A4).build(story)
